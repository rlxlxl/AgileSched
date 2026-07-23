#!/usr/bin/env python3
"""Generate docs/Руководство_пользователя_Расписание_РиМ.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Руководство_пользователя_Расписание_РиМ.docx"


def set_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    doc = Document()
    set_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Руководство пользователя")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Бот «Расписание РиМ»")
    r2.font.size = Pt(14)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Инструкция для сотрудников: как заполнять и смотреть своё расписание."
    )

    # 1
    add_heading(doc, "1. Что это", 1)
    add_para(
        doc,
        "Бот помогает заполнять общий Excel-файл «Расписашка РиМ». "
        "Файл лежит на Google Drive и синхронизируется у команды. "
        "Вы общаетесь с ботом в чате задачи YouGile (или в Telegram) — "
        "бот сам записывает часы в нужные ячейки.",
    )

    # 2
    add_heading(doc, "2. Как начать (YouGile)", 1)
    add_numbered(
        doc,
        [
            "Откройте любую задачу в YouGile.",
            "Справа откройте чат задачи.",
            "Напишите команду /start.",
            "Дальше отвечайте номерами пунктов меню (1, 2, 3…).",
        ],
    )
    add_para(
        doc,
        "После каждого действия бот снова показывает главное меню — "
        "повторно вводить /start обычно не нужно. "
        "Чтобы прервать текущий шаг, напишите /cancel.",
    )

    # 3
    add_heading(doc, "3. Главное меню", 1)
    add_bullets(
        doc,
        [
            "1 — Заполнить расписание",
            "2 — Привязать профиль",
            "3 — Моё расписание",
            "4 — Статус",
        ],
    )
    add_para(
        doc,
        "Можно писать и команды: /schedule, /my, /myschedule, /status, /cancel.",
    )

    # 4
    add_heading(doc, "4. Привязать профиль", 1)
    add_para(
        doc,
        "Сначала один раз привяжите себя к строке в Excel. Выберите пункт "
        "«Привязать профиль» (или /my) и укажите:",
    )
    add_bullets(
        doc,
        [
            "отдел;",
            "себя в списке сотрудников;",
            "ставку: 1 (полная, норма 40 ч/нед) или 0,5 (полставки, 20 ч/нед);",
            "обед: без обеда / 30 / 45 / 60 / 90 минут (вычитается из часов дня).",
        ],
    )
    add_para(
        doc,
        "После привязки при заполнении расписания бот сам подставит вашего "
        "сотрудника — отдел и ФИО выбирать заново не нужно. "
        "«Моё расписание» тоже использует этот профиль.",
    )

    # 5
    add_heading(doc, "5. Заполнить расписание", 1)
    add_para(
        doc,
        "Выберите пункт 1 «Заполнить расписание» (или /schedule).",
    )
    add_heading(doc, "Лист и неделя", 2)
    add_para(
        doc,
        "Бот сам выбирает лист Excel по сегодняшней дате "
        "(например, «Расписашка 2026 ИЮЛЬ») и текущую неделю на этом листе. "
        "Если не удалось определить — попросит выбрать лист или неделю вручную.",
    )
    add_heading(doc, "Как заполнить", 2)
    add_para(
        doc,
        "После выбора (или автоподстановки) сотрудника бот спросит способ:",
    )
    add_bullets(
        doc,
        [
            "Пошагово — дни, время и вид работы выбираются номерами из списков.",
            "Одним сообщением — вы пишете расписание текстом.",
        ],
    )
    add_para(doc, "Пример свободного ввода:", bold=True)
    add_para(doc, "Пн 9:00-18:00, Вт 10:00-19:00, Ср 9-18")
    add_para(
        doc,
        "Можно указывать несколько дней через запятую. "
        "Затем выберите вид работы и подтвердите сохранение.",
    )

    add_heading(doc, "Виды работы (что пишется в Excel)", 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ("Тип", "В ячейке", "Примечание")
    rows = (
        ("Полуочно (0,5 ч)", "0.5", "Полчаса слота"),
        ("Очно", "1", "Полный час, обычно зелёная заливка"),
        ("Дистанционно", "Д", "Считается как 1 час"),
    )
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_para(
        doc,
        "После сохранения бот покажет, сколько ячеек обновлено, и сводку часов "
        "за неделю (с учётом обеда и нормы).",
    )

    # 6
    add_heading(doc, "6. Моё расписание", 1)
    add_para(
        doc,
        "Пункт 3 «Моё расписание» (или /myschedule) показывает вашу текущую "
        "неделю: дни, часы, вычет обеда и сверку с нормой по ставке. "
        "Нужен привязанный профиль.",
    )

    # 7
    add_heading(doc, "7. Подзадачи и ответственный", 1)
    add_para(
        doc,
        "Если вы пишете боту в чате подзадачи, у которой назначен ответственный "
        "и у него уже есть профиль:",
    )
    add_bullets(
        doc,
        [
            "бот может автоматически подставить этого человека при заполнении;",
            "«Моё расписание» и «Статус» в таком чате относятся к ответственному;",
            "при смене ответственного бот может написать об этом в чат.",
        ],
    )
    add_para(
        doc,
        "Если у ответственного нет профиля — его нужно один раз привязать через «Привязать профиль».",
    )

    # 8
    add_heading(doc, "8. Как смотреть Excel", 1)
    add_bullets(
        doc,
        [
            "Открывайте файл как .xlsx в Excel / Numbers / скачивайте с Диска.",
            "Не открывайте через «Открыть в Google Таблицах» — цвета и стили часто ломаются.",
            "После сохранения ботом подождите 5–30 секунд, пока Drive синхронизирует файл.",
            "Не держите тот же файл открытым в Excel на компьютере, пока бот пишет — возможна блокировка.",
        ],
    )

    # 9
    add_heading(doc, "9. Telegram", 1)
    add_para(
        doc,
        "Тот же бот доступен в Telegram: /start открывает меню с кнопками. "
        "Действия те же — заполнить, профиль, моё расписание, статус. "
        "Запись идёт в тот же Excel-файл.",
    )

    # 10
    add_heading(doc, "10. Частые вопросы", 1)
    add_para(doc, "Бот снова просит выбрать отдел и сотрудника.", bold=True)
    add_para(
        doc,
        "Привяжите профиль (/my). После этого при заполнении человек подставится сам.",
    )
    add_para(doc, "Не тот месяц или неделя.", bold=True)
    add_para(
        doc,
        "Лист выбирается по сегодняшней дате (год и месяц в названии листа). "
        "Если автовыбор не сработал — выберите лист и неделю вручную из списка.",
    )
    add_para(doc, "В Excel нет изменений.", bold=True)
    add_para(
        doc,
        "Подождите синхронизацию Drive, обновите файл, убедитесь, что смотрите "
        "именно .xlsx, а не копию в Google Таблицах.",
    )
    add_para(doc, "Нужно отменить ввод.", bold=True)
    add_para(doc, "Напишите /cancel — бот вернёт в главное меню.")

    footer = doc.add_paragraph()
    footer.add_run(
        "\nПри проблемах с доступом к файлу или ошибках бота обратитесь "
        "к администратору YouGile / владельцу расширения «Расписание РиМ»."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
